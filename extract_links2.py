from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def extract_hyperlinks(doc):
    rels = doc.part.rels
    links = {}
    for rel_id, rel in rels.items():
        if rel.reltype == RT.HYPERLINK:
            links[rel_id] = rel.target_ref
            
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            pass # this is hard in python-docx
            
    # Better approach using xml:
    from docx.oxml.ns import qn
    for paragraph in doc.paragraphs:
        for child in paragraph._p:
            if child.tag == qn('w:hyperlink'):
                rel_id = child.get(qn('r:id'))
                if rel_id in links:
                    text = "".join(t.text for t in child.iter(qn('w:t')) if t.text)
                    print(f"[{text}]({links[rel_id]})")

extract_hyperlinks(Document("images/p3v/Article.docx"))
